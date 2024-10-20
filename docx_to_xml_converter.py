#!/usr/bin/env python3
"""
docx_to_xml_converter.py

A Python module to convert .docx files to XML format.
Enhanced to handle non-breaking spaces, section breaks, page breaks, and tables.
Designed with modularity, robustness, and extensibility in mind.

Author: OpenAI ChatGPT
Date: 2024-04-27
"""

import os
import sys
import logging
from typing import Any, Dict, List, Optional, Union
from docx import Document
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET

# Configure logging
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)  # Set to DEBUG for detailed logging

# Create console handler with a higher log level
ch = logging.StreamHandler()
ch.setLevel(logging.INFO)  # Change to DEBUG to see detailed logs on console

# Create file handler which logs even debug messages
fh = logging.FileHandler('docx_to_xml_converter.log')
fh.setLevel(logging.DEBUG)

# Create formatter and add it to the handlers
formatter = logging.Formatter(
    '%(asctime)s - %(name)s - %(levelname)s - %(message)s')
ch.setFormatter(formatter)
fh.setFormatter(formatter)

# Add the handlers to the logger
logger.addHandler(ch)
logger.addHandler(fh)


class DocxParser:
    """
    A class to parse .docx files and extract their contents, including
    paragraphs, tables, and special elements like non-breaking spaces,
    section breaks, and page breaks.
    """

    # Define the namespace mapping as a class variable
    NAMESPACE = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    def __init__(self, filepath: str):
        """
        Initializes the DocxParser with the path to the .docx file.

        :param filepath: Path to the .docx file to be parsed.
        """
        self.filepath = filepath
        self.document = None
        logger.debug(f"DocxParser initialized with file: {self.filepath}")

    def load_document(self) -> None:
        """
        Loads the .docx document.

        :raises FileNotFoundError: If the file does not exist.
        :raises Exception: If the file cannot be opened as a .docx.
        """
        logger.debug("Attempting to load the .docx document.")
        if not os.path.isfile(self.filepath):
            logger.error(f"File not found: {self.filepath}")
            raise FileNotFoundError(f"File not found: {self.filepath}")
        try:
            self.document = Document(self.filepath)
            logger.info(f"Successfully loaded document: {self.filepath}")
        except Exception as e:
            logger.error(f"Error loading document: {e}")
            raise

    def parse_document(self) -> Dict[str, Any]:
        """
        Parses the loaded document and extracts its content, including
        paragraphs and tables.

        :return: A dictionary containing parsed paragraphs and tables.
        :raises Exception: If the document is not loaded or cannot be parsed.
        """
        logger.debug("Starting to parse the document.")
        if self.document is None:
            logger.error("Document not loaded. Call load_document() first.")
            raise Exception("Document not loaded. Call load_document() first.")

        parsed_content = {
            'paragraphs': self.parse_paragraphs(),
            'tables': self.parse_tables()
        }
        logger.info("Document parsing completed successfully.")
        return parsed_content

    def parse_paragraphs(self) -> List[Dict[str, Any]]:
        """
        Parses the paragraphs in the document, detecting non-breaking spaces,
        section breaks, and page breaks.

        :return: A list of dictionaries containing parsed paragraph content.
        """
        logger.debug("Parsing paragraphs.")
        parsed_paragraphs = []
        for para in self.document.paragraphs:
            para_dict = {
                'text_elements': self.parse_text_runs(para),
                'style': para.style.name,
                'breaks': self.detect_paragraph_breaks(para)
            }
            parsed_paragraphs.append(para_dict)
            logger.debug(f"Parsed paragraph: {para_dict}")
        return parsed_paragraphs

    def parse_text_runs(self, para) -> List[Dict[str, Union[str, bool]]]:
        """
        Parses the runs within a paragraph to detect formatting and non-breaking spaces.

        :param para: A paragraph object from python-docx.
        :return: A list of dictionaries representing text elements.
        """
        text_elements = []
        for run in para.runs:
            # Detect non-breaking spaces
            text = run.text
            if '\u00A0' in text:
                parts = text.split('\u00A0')
                for i, part in enumerate(parts):
                    if part:
                        text_elements.append({
                            'text': part,
                            'bold': run.bold or False,
                            'italic': run.italic or False,
                            'underline': run.underline or False,
                            'non_breaking_space': False
                        })
                    if i < len(parts) - 1:
                        text_elements.append({
                            'text': '',
                            'bold': False,
                            'italic': False,
                            'underline': False,
                            'non_breaking_space': True
                        })
            else:
                text_elements.append({
                    'text': text,
                    'bold': run.bold or False,
                    'italic': run.italic or False,
                    'underline': run.underline or False,
                    'non_breaking_space': False
                })
        return text_elements

    def detect_paragraph_breaks(self, para) -> List[str]:
        """
        Detects section breaks and page breaks within a paragraph.

        :param para: A paragraph object from python-docx.
        :return: A list of break types detected in the paragraph.
        """
        breaks = []
        for run in para.runs:
            # Access the underlying XML element using _element
            for br in run._element.findall('.//w:br', namespaces=self.NAMESPACE):
                br_type = br.get(qn('w:type'))
                if br_type == 'page':
                    breaks.append('page_break')
                elif br_type in ('column', 'text_wrapping'):
                    # Additional break types can be handled here
                    pass
                else:
                    # Default or unknown break types
                    breaks.append('section_break')
        return breaks

    def parse_tables(self) -> List[Dict[str, Any]]:
        """
        Parses the tables in the document.

        :return: A list of dictionaries containing parsed table data.
        """
        logger.debug("Parsing tables.")
        parsed_tables = []
        for table in self.document.tables:
            table_dict = {
                'rows': []
            }
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Concatenate all paragraph texts within the cell, separated by newlines
                    cell_text = '\n'.join([para.text for para in cell.paragraphs])
                    row_data.append(cell_text)
                table_dict['rows'].append(row_data)
            parsed_tables.append(table_dict)
            logger.debug(f"Parsed table: {table_dict}")
        return parsed_tables


class XMLConverter:
    """
    A class to convert parsed document data into an XML structure,
    including paragraphs, non-breaking spaces, breaks, and tables.
    """

    def __init__(self, data: Dict[str, Any]):
        """
        Initializes the XMLConverter with parsed data.

        :param data: Parsed data from the .docx file.
        """
        self.data = data
        logger.debug("XMLConverter initialized with parsed data.")

    def build_xml(self) -> ET.Element:
        """
        Builds the XML ElementTree from the parsed data.

        :return: The root XML element.
        """
        logger.debug("Starting to build XML structure.")
        root = ET.Element('Document')

        # Add Paragraphs
        paragraphs_element = ET.SubElement(root, 'Paragraphs')
        for idx, para in enumerate(self.data.get('paragraphs', []), start=1):
            para_element = ET.SubElement(paragraphs_element, 'Paragraph', id=str(idx))

            # Add Style
            style_element = ET.SubElement(para_element, 'Style')
            style_element.text = para['style']

            # Add Text Elements
            text_elements_element = ET.SubElement(para_element, 'TextElements')
            for text_idx, text_element in enumerate(para['text_elements'], start=1):
                if text_element['non_breaking_space']:
                    nbsp_element = ET.SubElement(text_elements_element, 'NonBreakingSpace', id=str(text_idx))
                    nbsp_element.text = ''
                else:
                    run_attributes = {
                        'bold': str(text_element['bold']),
                        'italic': str(text_element['italic']),
                        'underline': str(text_element['underline'])
                    }
                    run_element = ET.SubElement(text_elements_element, 'Run', **run_attributes)
                    run_element.text = text_element['text']

            # Add Breaks if any
            if para['breaks']:
                breaks_element = ET.SubElement(para_element, 'Breaks')
                for break_type in para['breaks']:
                    # Capitalize the first letter for XML element naming
                    break_element = ET.SubElement(breaks_element, f"{break_type.capitalize()}")
                    break_element.text = 'True'

            logger.debug(f"Added XML for paragraph {idx}: {para}")

        # Add Tables
        tables_element = ET.SubElement(root, 'Tables')
        for tbl_idx, table in enumerate(self.data.get('tables', []), start=1):
            table_element = ET.SubElement(tables_element, 'Table', id=str(tbl_idx))
            for row_idx, row in enumerate(table['rows'], start=1):
                row_element = ET.SubElement(table_element, 'Row', id=str(row_idx))
                for cell_idx, cell in enumerate(row, start=1):
                    cell_element = ET.SubElement(row_element, 'Cell', id=str(cell_idx))
                    cell_element.text = cell
            logger.debug(f"Added XML for table {tbl_idx}: {table}")

        logger.info("XML structure built successfully.")
        return root

    def save_xml(self, root: ET.Element, output_path: str) -> None:
        """
        Saves the XML ElementTree to a file.

        :param root: The root XML element.
        :param output_path: Path where the XML file will be saved.
        :raises Exception: If the file cannot be written.
        """
        logger.debug(f"Attempting to save XML to: {output_path}")
        try:
            tree = ET.ElementTree(root)
            tree.write(output_path, encoding='utf-8', xml_declaration=True)
            logger.info(f"XML file saved successfully at: {output_path}")
        except Exception as e:
            logger.error(f"Error saving XML file: {e}")
            raise


class Converter:
    """
    A facade class to convert a .docx file to an XML file, handling parsing and conversion.
    """

    def __init__(self, input_path: str, output_path: Optional[str] = None):
        """
        Initializes the Converter with input and output paths.

        :param input_path: Path to the input .docx file.
        :param output_path: Path to the output XML file. If None, replaces .docx with .xml.
        """
        self.input_path = input_path
        if output_path:
            self.output_path = output_path
        else:
            base, _ = os.path.splitext(input_path)
            self.output_path = f"{base}.xml"
        logger.debug(f"Converter initialized with input: {self.input_path}, output: {self.output_path}")

    def convert(self) -> None:
        """
        Performs the conversion from .docx to XML.

        :raises Exception: If any step in the conversion fails.
        """
        logger.info(f"Starting conversion: {self.input_path} -> {self.output_path}")
        parser = DocxParser(self.input_path)
        parser.load_document()
        parsed_data = parser.parse_document()

        converter = XMLConverter(parsed_data)
        xml_root = converter.build_xml()
        converter.save_xml(xml_root, self.output_path)
        logger.info("Conversion completed successfully.")


def convert_docx_to_xml(input_path: str, output_path: Optional[str] = None) -> None:
    """
    Converts a .docx file to an XML file.

    :param input_path: Path to the input .docx file.
    :param output_path: Path to the output XML file. If None, replaces .docx with .xml.
    """
    try:
        converter = Converter(input_path, output_path)
        converter.convert()
    except Exception as e:
        logger.exception(f"Failed to convert {input_path} to XML. Error: {e}")
        raise


def main():
    """
    The main function to execute when running the script directly.
    Parses command-line arguments and performs the conversion.
    """
    import argparse

    parser = argparse.ArgumentParser(description='Convert a .docx file to XML format.')
    parser.add_argument('input', help='Path to the input .docx file.')
    parser.add_argument('-o', '--output', help='Path to the output XML file.')

    args = parser.parse_args()

    try:
        convert_docx_to_xml(args.input, args.output)
    except Exception as e:
        logger.error(f"An error occurred during conversion: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
